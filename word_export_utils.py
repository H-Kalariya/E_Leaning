from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os


# ─── LaTeX → Unicode / plain text conversion ──────────────────────────────────

SYMBOL_MAP = [
    # Greek letters
    (r'\\psi',      'ψ'), (r'\\Psi',      'Ψ'),
    (r'\\phi',      'ϕ'), (r'\\Phi',      'Φ'),
    (r'\\theta',    'θ'), (r'\\Theta',    'Θ'),
    (r'\\pi',       'π'), (r'\\Pi',       'Π'),
    (r'\\alpha',    'α'), (r'\\Alpha',    'Α'),
    (r'\\beta',     'β'), (r'\\Beta',     'Β'),
    (r'\\gamma',    'γ'), (r'\\Gamma',    'Γ'),
    (r'\\delta',    'δ'), (r'\\Delta',    'Δ'),
    (r'\\epsilon',  'ε'), (r'\\Epsilon',  'Ε'),
    (r'\\zeta',     'ζ'), (r'\\eta',      'η'),
    (r'\\iota',     'ι'), (r'\\kappa',    'κ'),
    (r'\\lambda',   'λ'), (r'\\Lambda',   'Λ'),
    (r'\\mu',       'μ'), (r'\\nu',       'ν'),
    (r'\\xi',       'ξ'), (r'\\Xi',       'Ξ'),
    (r'\\rho',      'ρ'), (r'\\sigma',    'σ'),
    (r'\\Sigma',    'Σ'), (r'\\tau',      'τ'),
    (r'\\upsilon',  'υ'), (r'\\chi',      'χ'),
    (r'\\omega',    'ω'), (r'\\Omega',    'Ω'),
    (r'\\hbar',     'ℏ'),
    # Operators & relations
    (r'\\infty',    '∞'), (r'\\sqrt',     '√'),
    (r'\\int',      '∫'), (r'\\oint',     '∮'),
    (r'\\sum',      '∑'), (r'\\prod',     '∏'),
    (r'\\partial',  '∂'), (r'\\nabla',    '∇'),
    (r'\\approx',   '≈'), (r'\\neq',      '≠'),
    (r'\\equiv',    '≡'), (r'\\leq',      '≤'),
    (r'\\geq',      '≥'), (r'\\le\b',     '≤'),
    (r'\\ge\b',     '≥'), (r'\\ll',       '≪'),
    (r'\\gg',       '≫'), (r'\\times',    '×'),
    (r'\\cdot',     '·'), (r'\\pm',       '±'),
    (r'\\mp',       '∓'), (r'\\div',      '÷'),
    (r'\\circ',     '∘'), (r'\\propto',   '∝'),
    (r'\\sim',      '~'), (r'\\simeq',    '≃'),
    # Arrows
    (r'\\rightarrow',  '→'), (r'\\leftarrow',   '←'),
    (r'\\Rightarrow',  '⇒'), (r'\\Leftarrow',   '⇐'),
    (r'\\leftrightarrow', '↔'), (r'\\Leftrightarrow', '⇔'),
    (r'\\uparrow',     '↑'), (r'\\downarrow',    '↓'),
    # Misc
    (r'\\dots',     '…'), (r'\\ldots',    '…'),
    (r'\\cdots',    '⋯'), (r'\\forall',   '∀'),
    (r'\\exists',   '∃'), (r'\\in\b',     '∈'),
    (r'\\notin',    '∉'), (r'\\subset',   '⊂'),
    (r'\\supset',   '⊃'), (r'\\cup',      '∪'),
    (r'\\cap',      '∩'), (r'\\emptyset', '∅'),
    (r'\\angle',    '∠'), (r'\\perp',     '⊥'),
    (r'\\parallel', '∥'),
]

def clean_math(expr):
    """
    Convert a LaTeX math expression to a readable plain-text / Unicode form
    suitable for Word documents.
    """
    # Remove wrapping $ or $$ delimiters (if any accidentally left in)
    expr = re.sub(r'^\$\$?|\$\$?$', '', expr.strip())

    # \frac{a}{b} → (a/b)
    expr = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'(\1/\2)', expr)
    # \sqrt{x} → √(x)
    expr = re.sub(r'\\sqrt\{([^}]*)\}', r'√(\1)', expr)
    # \text{...} → content
    expr = re.sub(r'\\text\{([^}]*)\}', r'\1', expr)
    # \mathrm{...} / \mathbf{...} etc → content
    expr = re.sub(r'\\math(?:rm|bf|it|sf|tt|cal|bb)\{([^}]*)\}', r'\1', expr)
    # \left( \right) etc → ( )
    expr = re.sub(r'\\(?:left|right)[.|()\[\]{}|]', '', expr)
    # \overline{x} → x̄  (approximate)
    expr = re.sub(r'\\overline\{([^}]*)\}', r'\1̄', expr)
    # ^{...} → superscript in text form
    expr = re.sub(r'\^\{([^}]*)\}', r'^\1', expr)
    # _{...} → subscript in text form
    expr = re.sub(r'_\{([^}]*)\}', r'_\1', expr)
    # Simple ^ and _
    expr = re.sub(r'\^([^\s{])', r'^\1', expr)
    expr = re.sub(r'_([^\s{])', r'_\1', expr)

    # Apply symbol map
    for pattern, replacement in SYMBOL_MAP:
        expr = re.sub(pattern, replacement, expr)

    # Strip remaining backslash commands we couldn't convert
    expr = re.sub(r'\\[a-zA-Z]+', '', expr)
    # Clean up braces
    expr = expr.replace('{', '').replace('}', '')
    return expr.strip()


def render_latex_in_text(text):
    """
    Replace all LaTeX math regions in a line of Markdown text with their
    Unicode / plain-text equivalents, returning the cleaned string.
    Also handles \\( \\) and \\[ \\] delimiter variants.
    """
    # Normalise \( \) and \[ \] → $ and $$
    text = re.sub(r'\\\[\s*', '$$', text)
    text = re.sub(r'\s*\\\]', '$$', text)
    text = re.sub(r'\\\(', '$', text)
    text = re.sub(r'\\\)', '$', text)

    # Block math $$...$$
    text = re.sub(r'\$\$(.+?)\$\$', lambda m: clean_math(m.group(1)), text, flags=re.DOTALL)
    # Inline math $...$
    text = re.sub(r'\$([^$\n]+?)\$', lambda m: clean_math(m.group(1)), text)
    return text


# ─── Inline markdown formatter for a paragraph ────────────────────────────────

def add_formatted_run(paragraph, text):
    """
    Parse inline markdown (**bold**, *italic*, `code`) and add styled runs
    to a python-docx paragraph object.
    """
    # Split on bold, italic-bold, italic, or code
    token_re = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`[^`]+`)')
    parts = token_re.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


# ─── Main export function ─────────────────────────────────────────────────────

def create_word_document(content, output_path):
    """
    Creates a professionally formatted Word document from Markdown + LaTeX content.

    Supports:
      - # / ## / ### headings  (Word Heading 1/2/3)
      - **bold**, *italic*, `code` inline formatting
      - Bullet lists  - item  /  * item
      - Numbered lists  1. item
      - Code blocks  ```...```
      - Display math  $$...$$  (rendered as styled paragraph)
      - Inline math  $...$  (rendered inline)
      - Blank lines preserved as paragraph spacing
    """
    doc = Document()

    # ── Document-wide defaults ────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    # Paragraph spacing
    style.paragraph_format.space_after = Pt(6)

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # ── Code block handling ───────────────────────────────────────────────
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                i += 1
                continue
            else:
                # End of code block — emit it
                in_code_block = False
                if code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.left_indent = Inches(0.3)
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x1f, 0x26, 0x35)
                i += 1
                continue

        if in_code_block:
            code_lines.append(raw)
            i += 1
            continue

        # ── Blank line → small spacer ─────────────────────────────────────────
        if not line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # ── Display math  $$...$$ on its own line ─────────────────────────────
        if line.strip().startswith('$$') and line.strip().endswith('$$') and len(line.strip()) > 4:
            expr = line.strip()[2:-2].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(clean_math(expr))
            run.bold = True
            run.font.name = 'Cambria Math'
            run.font.size = Pt(12)
            i += 1
            continue

        # ── Standalone $$ block opener (content on next lines) ────────────────
        if line.strip() == '$$':
            math_lines = []
            i += 1
            while i < len(lines) and lines[i].rstrip() != '$$':
                math_lines.append(lines[i].rstrip())
                i += 1
            expr = '\n'.join(math_lines)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(clean_math(expr))
            run.bold = True
            run.font.name = 'Cambria Math'
            run.font.size = Pt(12)
            i += 1  # skip closing $$
            continue

        # ── Headings ──────────────────────────────────────────────────────────
        if line.startswith('#### '):
            p = doc.add_heading(render_latex_in_text(line[5:]), level=3)
            p.paragraph_format.space_before = Pt(10)
            i += 1
            continue
        if line.startswith('### '):
            p = doc.add_heading(render_latex_in_text(line[4:]), level=3)
            p.paragraph_format.space_before = Pt(12)
            i += 1
            continue
        if line.startswith('## '):
            p = doc.add_heading(render_latex_in_text(line[3:]), level=2)
            p.paragraph_format.space_before = Pt(14)
            i += 1
            continue
        if line.startswith('# '):
            p = doc.add_heading(render_latex_in_text(line[2:]), level=1)
            p.paragraph_format.space_before = Pt(16)
            i += 1
            continue

        # ── Horizontal rule ────────────────────────────────────────────────────
        if re.match(r'^[-*_]{3,}$', line.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('─' * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            i += 1
            continue

        # ── Bullet list  - item  or  * item ───────────────────────────────────
        bullet_match = re.match(r'^(\s*)[*\-+] (.+)', line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            item_text = render_latex_in_text(bullet_match.group(2))
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * (1 + indent // 2))
            p.paragraph_format.space_after = Pt(3)
            add_formatted_run(p, item_text)
            i += 1
            continue

        # ── Numbered list  1. / 2. etc ─────────────────────────────────────────
        num_match = re.match(r'^\s*\d+\. (.+)', line)
        if num_match:
            item_text = render_latex_in_text(num_match.group(1))
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            add_formatted_run(p, item_text)
            i += 1
            continue

        # ── Blockquote  > text ──────────────────────────────────────────────────
        if line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(render_latex_in_text(line[2:]))
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # ── Normal paragraph ────────────────────────────────────────────────────
        clean = render_latex_in_text(line)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        add_formatted_run(p, clean)
        i += 1

    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    test_content = """# Quantum Mechanics: Wave Functions

## 1. Introduction

**Quantum mechanics** (QM) describes nature at the smallest scales of energy levels of *atoms* and *subatomic particles*.

## 2. The Wave Function

The state of a quantum system is described by a **wave function** $\\psi(x, t)$.

The probability of finding a particle between $x$ and $x + dx$ is:

$$|\\psi(x,t)|^2 \\, dx$$

### 2.1 Normalisation Condition

For any valid wave function:

$$\\int_{-\\infty}^{\\infty} |\\psi(x,t)|^2 \\, dx = 1$$

### 2.2 Schrödinger Equation

The time-dependent Schrödinger equation is:

$$i\\hbar \\frac{\\partial}{\\partial t}\\psi = \\hat{H}\\psi$$

where $\\hat{H}$ is the **Hamiltonian operator**.

## 3. Key Principles

- **Superposition**: A system can exist in multiple states simultaneously.
- **Entanglement**: Two particles can be correlated regardless of distance.
- **Uncertainty**: $\\Delta x \\cdot \\Delta p \\geq \\frac{\\hbar}{2}$

## 4. Example Code (Python)

```python
import numpy as np

# Simple harmonic oscillator ground state
def psi(x, a=1.0):
    return (a / np.pi)**0.25 * np.exp(-a * x**2 / 2)
```

## 5. Summary

Quantum mechanics revolutionised our understanding of the microscopic world.
The core quantity, $\\psi$, encodes all information about the system.
"""
    out = 'test_quantum_notes.docx'
    create_word_document(test_content, out)
    print(f'✅ Test document created: {os.path.abspath(out)}')
